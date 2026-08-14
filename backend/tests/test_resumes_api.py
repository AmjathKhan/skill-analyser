"""Resume upload, download, duplicate detection and reprocessing."""

from __future__ import annotations

from pathlib import Path

from fastapi.testclient import TestClient

SAMPLES = Path(__file__).resolve().parents[1] / "data" / "sample_resumes"


def test_upload_parses_and_creates_candidates(uploaded_candidates) -> None:
    assert len(uploaded_candidates) >= 3
    for item in uploaded_candidates:
        assert item["status"] == "completed", item
        assert item["candidate_id"], item
        processing = item["processing"]
        assert processing["skills_normalized"] > 0
        assert processing["embeddings"] > 0
        assert processing["graph_edges"] > 0
        assert processing["error"] is None


def test_duplicate_upload_is_detected(client: TestClient, admin_headers, uploaded_candidates) -> None:
    path = SAMPLES / "priya_sharma_fullstack.txt"
    response = client.post(
        "/api/upload",
        files=[("files", (path.name, path.read_bytes(), "text/plain"))],
        headers=admin_headers,
    )
    assert response.status_code == 201
    body = response.json()
    assert body["duplicates"] == 1
    assert body["uploaded"] == 0
    assert body["results"][0]["is_duplicate"] is True
    assert body["results"][0]["duplicate_of_resume_id"]


def test_unsupported_extension_is_rejected(client: TestClient, admin_headers) -> None:
    response = client.post(
        "/api/upload",
        files=[("files", ("malware.exe", b"MZ\x90\x00binary", "application/octet-stream"))],
        headers=admin_headers,
    )
    assert response.status_code == 201
    body = response.json()
    assert body["failed"] == 1
    assert body["results"][0]["error"]


def test_pdf_upload_is_parsed(client: TestClient, admin_headers) -> None:
    import io

    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer)
    for index, line in enumerate(
        [
            "Vikram Rao",
            "vikram.rao@example.com | +91 90000 12345",
            "SUMMARY",
            "Backend engineer with 7 years of experience.",
            "TECHNICAL SKILLS",
            "Python, FastAPI, PostgreSQL, Docker, AWS, Redis",
            "EXPERIENCE",
            "Senior Backend Engineer - Helios Systems (Jan 2020 - Present)",
            "EDUCATION",
            "B.E. Computer Science, Anna University, 2016",
        ]
    ):
        pdf.drawString(60, 780 - index * 18, line)
    pdf.save()

    response = client.post(
        "/api/upload?wait=true",
        files=[("files", ("vikram_rao.pdf", buffer.getvalue(), "application/pdf"))],
        headers=admin_headers,
    )
    assert response.status_code == 201, response.text
    result = response.json()["results"][0]
    assert result["status"] == "completed", result
    candidate_id = result["candidate_id"]

    detail = client.get(f"/api/candidate/{candidate_id}", headers=admin_headers).json()
    assert detail["full_name"] == "Vikram Rao"
    assert detail["email"] == "vikram.rao@example.com"
    skill_names = {skill["name"].lower() for skill in detail["skills"]}
    assert {"python", "fastapi", "postgresql"} <= skill_names


def _build_docx_resume() -> bytes:
    from io import BytesIO

    from docx import Document

    buffer = BytesIO()
    document = Document()
    document.add_paragraph("Ananya Reddy")
    document.add_paragraph("ananya.reddy@example.com | +91 98888 11223")
    document.add_heading("Summary", level=2)
    document.add_paragraph("Full stack engineer with 6 years of experience.")
    document.add_heading("Technical Skills", level=2)
    document.add_paragraph("Python, FastAPI, React, PostgreSQL, Docker, AWS")
    document.add_heading("Experience", level=2)
    document.add_paragraph("Senior Software Engineer - Helios Systems (Mar 2019 - Present)")
    document.add_heading("Education", level=2)
    document.add_paragraph("B.Tech Computer Science, NIT Warangal, 2018")
    document.save(buffer)
    return buffer.getvalue()


def test_docx_upload_is_parsed(client: TestClient, admin_headers) -> None:
    response = client.post(
        "/api/upload?wait=true",
        files=[
            (
                "files",
                (
                    "ananya_reddy.docx",
                    _build_docx_resume(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            )
        ],
        headers=admin_headers,
    )
    assert response.status_code == 201, response.text
    result = response.json()["results"][0]
    assert result["status"] == "completed", result
    candidate_id = result["candidate_id"]

    detail = client.get(f"/api/candidate/{candidate_id}", headers=admin_headers).json()
    assert detail["full_name"] == "Ananya Reddy"
    assert detail["email"] == "ananya.reddy@example.com"
    skill_names = {skill["name"].lower() for skill in detail["skills"]}
    assert {"python", "fastapi", "postgresql"} <= skill_names


def test_pdf_with_leading_bytes_is_accepted(client: TestClient, admin_headers) -> None:
    import io

    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.drawString(60, 760, "Kiran Das")
    pdf.drawString(60, 740, "kiran.das@example.com")
    pdf.drawString(60, 720, "TECHNICAL SKILLS")
    pdf.drawString(60, 700, "Python, Docker, AWS")
    pdf.drawString(60, 680, "EXPERIENCE")
    pdf.drawString(60, 660, "DevOps Engineer - Helios Systems (Jan 2021 - Present)")
    pdf.save()
    payload = b"\n" + buffer.getvalue()

    response = client.post(
        "/api/upload?wait=true",
        files=[("files", ("kiran_das.pdf", payload, "application/pdf"))],
        headers=admin_headers,
    )
    assert response.status_code == 201, response.text
    result = response.json()["results"][0]
    assert result["status"] == "completed", result
    assert result["candidate_id"]


def test_list_download_and_status(client: TestClient, admin_headers, uploaded_candidates) -> None:
    listing = client.get("/api/resumes?page_size=50", headers=admin_headers)
    assert listing.status_code == 200
    items = listing.json()["items"]
    assert items
    resume_id = items[0]["id"]

    detail = client.get(f"/api/resume/{resume_id}", headers=admin_headers)
    assert detail.status_code == 200
    assert detail.json()["raw_text"]
    assert detail.json()["parsed_data"]["personal"]["full_name"]

    status = client.get(f"/api/resume/{resume_id}/status", headers=admin_headers)
    assert status.status_code == 200
    assert status.json()["status"] in {"completed", "duplicate"}

    download = client.get(f"/api/resume/{resume_id}/download", headers=admin_headers)
    assert download.status_code == 200
    assert download.content
    assert "attachment" in download.headers["content-disposition"]

    inline = client.get(f"/api/resume/{resume_id}/download?inline=true", headers=admin_headers)
    assert "inline" in inline.headers["content-disposition"]


def test_reprocess_is_idempotent(client: TestClient, admin_headers, uploaded_candidates) -> None:
    resume_id = next(item["resume_id"] for item in uploaded_candidates)
    before = client.get(f"/api/candidate/{uploaded_candidates[0]['candidate_id']}", headers=admin_headers).json()

    response = client.post(f"/api/resume/{resume_id}/reprocess", headers=admin_headers)
    assert response.status_code == 200
    assert response.json()["status"] == "completed"

    after = client.get(f"/api/candidate/{uploaded_candidates[0]['candidate_id']}", headers=admin_headers).json()
    assert len(after["skills"]) == len(before["skills"])


def test_missing_resume_returns_404(client: TestClient, admin_headers) -> None:
    response = client.get("/api/resume/999999", headers=admin_headers)
    assert response.status_code == 404
    assert response.json()["error"]["code"] == "not_found"
