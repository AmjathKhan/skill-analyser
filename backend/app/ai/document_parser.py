"""Text extraction from PDF / DOCX / DOC / TXT resumes, with optional OCR."""

from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from xml.etree import ElementTree as ET

from app.ai.text_utils import clean_text
from app.core.config import settings
from app.core.exceptions import UnsupportedFileError

#: Below this many characters a PDF is assumed to be a scan and OCR is attempted.
OCR_TEXT_THRESHOLD = 220
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_DOCX_XML_PARTS = ("word/document.xml", "word/header", "word/footer", "word/footnotes", "word/endnotes")


@dataclass(slots=True)
class DocumentContent:
    text: str
    page_count: int = 0
    word_count: int = 0
    backend: str = "unknown"
    ocr_used: bool = False
    warnings: list[str] = field(default_factory=list)

    @property
    def is_empty(self) -> bool:
        return len(self.text.strip()) < 30


def _text_quality(text: str) -> int:
    """Prefer extracts that look like a resume (words + contact signals)."""
    stripped = (text or "").strip()
    if not stripped:
        return 0
    score = len(stripped.split())
    lowered = stripped.lower()
    if "@" in stripped:
        score += 80
    if "linkedin" in lowered or "github.com" in lowered:
        score += 20
    if any(token in lowered for token in ("experience", "education", "skills", "project")):
        score += 30
    return score


def sniff_document_suffix(path: Path, named_suffix: str | None = None) -> str:
    """Return the real document type from magic bytes, falling back to the filename."""
    named = (named_suffix or path.suffix).lower()
    header = path.read_bytes()[:2048]
    if b"%PDF" in header[:1024]:
        return ".pdf"
    if header.startswith(b"PK\x03\x04"):
        return ".docx"
    if header.startswith(b"\xd0\xcf\x11\xe0") or header.lstrip().startswith(b"{\\rtf"):
        return ".doc"
    return named


def extract_document(path: str | Path) -> DocumentContent:
    """Extract plain text from a resume file, choosing the parser by content then extension."""
    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"Resume file not found: {file_path}")

    named = file_path.suffix.lower()
    suffix = sniff_document_suffix(file_path, named)
    if named in {".txt", ".md"} and suffix not in {".pdf", ".docx", ".doc"}:
        suffix = named

    if suffix == ".pdf":
        content = _extract_pdf(file_path)
    elif suffix == ".docx":
        content = _extract_docx(file_path)
    elif suffix == ".doc":
        content = _extract_doc(file_path)
    elif suffix in {".txt", ".md"}:
        content = DocumentContent(text=file_path.read_text(encoding="utf-8", errors="ignore"), backend="plaintext")
    else:
        raise UnsupportedFileError(f"Unsupported resume format: {suffix or 'unknown'}")

    content.text = clean_text(content.text)
    content.word_count = len(content.text.split())
    if content.page_count == 0 and content.word_count:
        content.page_count = max(1, round(content.word_count / 500))
    return content


def extract_from_bytes(data: bytes, filename: str) -> DocumentContent:
    """Convenience wrapper used by tests and by in-memory processing."""
    suffix = Path(filename).suffix.lower() or ".txt"
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as handle:
        handle.write(data)
        handle.flush()
        temp_path = Path(handle.name)
    try:
        return extract_document(temp_path)
    finally:
        temp_path.unlink(missing_ok=True)


def _best_extract(*candidates: tuple[str, int, str]) -> tuple[str, int, str]:
    viable = [(text, pages, backend) for text, pages, backend in candidates if (text or "").strip()]
    if not viable:
        return "", 0, "none"
    return max(viable, key=lambda item: _text_quality(item[0]))


# --------------------------------------------------------------------------- PDF
def _extract_pdf(path: Path) -> DocumentContent:
    warnings: list[str] = []
    plumber_text, plumber_pages = _extract_pdf_plumber(path, warnings)
    fitz_text, fitz_pages = _extract_pdf_pymupdf(path, warnings)
    text, pages, backend = _best_extract(
        (plumber_text, plumber_pages, "pdfplumber"),
        (fitz_text, fitz_pages, "pymupdf"),
    )

    ocr_used = False
    if len(text.strip()) < OCR_TEXT_THRESHOLD and settings.enable_ocr:
        ocr_text = _ocr_pdf(path, warnings)
        if _text_quality(ocr_text) > _text_quality(text):
            text, backend, ocr_used = ocr_text, "tesseract-ocr", True
    elif len(text.strip()) < OCR_TEXT_THRESHOLD:
        warnings.append("Very little text extracted - the PDF may be scanned. Enable OCR with ENABLE_OCR=true.")

    return DocumentContent(text=text, page_count=pages, backend=backend, ocr_used=ocr_used, warnings=warnings)


def _extract_pdf_plumber(path: Path, warnings: list[str]) -> tuple[str, int]:
    try:
        import pdfplumber

        plain_parts: list[str] = []
        layout_parts: list[str] = []
        pages = 0
        with pdfplumber.open(str(path)) as pdf:
            pages = len(pdf.pages)
            for page in pdf.pages:
                plain = page.extract_text() or ""
                try:
                    layout = page.extract_text(layout=True) or ""
                except Exception:
                    layout = plain
                extra: list[str] = []
                try:
                    for table in page.extract_tables() or []:
                        for row in table:
                            cells = [cell for cell in row if cell]
                            if cells:
                                extra.append(" | ".join(cells))
                except Exception:
                    pass
                suffix = ("\n" + "\n".join(extra)) if extra else ""
                plain_parts.append(plain + suffix)
                layout_parts.append(layout + suffix)
        plain = "\n".join(plain_parts)
        layout = "\n".join(layout_parts)
        return (layout, pages) if _text_quality(layout) >= _text_quality(plain) else (plain, pages)
    except Exception as exc:
        warnings.append(f"pdfplumber failed: {exc.__class__.__name__}")
        return "", 0


def _extract_pdf_pymupdf(path: Path, warnings: list[str]) -> tuple[str, int]:
    try:
        import fitz  # PyMuPDF

        with fitz.open(str(path)) as doc:
            text = "\n".join(page.get_text("text") or "" for page in doc)
            blocks: list[str] = []
            for page in doc:
                for block in page.get_text("blocks") or []:
                    if len(block) >= 7 and block[6] == 0 and block[4]:
                        blocks.append(str(block[4]))
            blocks_text = "\n".join(blocks)
            chosen = blocks_text if _text_quality(blocks_text) > _text_quality(text) else text
            return chosen, doc.page_count
    except Exception as exc:
        warnings.append(f"pymupdf failed: {exc.__class__.__name__}")
        return "", 0


def _ocr_pdf(path: Path, warnings: list[str]) -> str:
    try:  # pragma: no cover - requires tesseract + AI extras
        import fitz
        import pytesseract
        from PIL import Image

        if settings.tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd

        parts: list[str] = []
        with fitz.open(str(path)) as doc:
            for page in doc:
                pixmap = page.get_pixmap(dpi=220)
                image = Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples)
                parts.append(pytesseract.image_to_string(image))
        return "\n".join(parts)
    except Exception as exc:
        warnings.append(f"OCR failed: {exc}")
        return ""


# -------------------------------------------------------------------------- DOCX
def _extract_docx(path: Path) -> DocumentContent:
    warnings: list[str] = []
    python_text = ""
    try:
        python_text = _extract_docx_python_docx(path)
    except Exception as exc:
        warnings.append(f"python-docx failed: {exc.__class__.__name__}")

    xml_text = ""
    try:
        xml_text = _extract_docx_xml(path)
    except Exception as exc:
        warnings.append(f"docx-xml failed: {exc.__class__.__name__}")

    text, _, backend = _best_extract(
        (python_text, 0, "python-docx"),
        (xml_text, 0, "docx-xml"),
    )
    return DocumentContent(text=text, backend=backend, warnings=warnings)


def _extract_docx_python_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        _collect_table_text(table, parts)
    for section in document.sections:
        for container in (section.header, section.footer):
            parts.extend(paragraph.text for paragraph in container.paragraphs if paragraph.text.strip())
    return "\n".join(part for part in parts if part and part.strip())


def _collect_table_text(table, parts: list[str]) -> None:
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
        if cells:
            parts.append(" | ".join(cells))
        for cell in row.cells:
            for nested in cell.tables:
                _collect_table_text(nested, parts)


def _extract_docx_xml(path: Path) -> str:
    """Read every ``w:t`` node, including text boxes and content controls python-docx skips."""
    parts: list[str] = []
    with zipfile.ZipFile(path) as archive:
        names = [
            name
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml") and any(name.startswith(prefix) for prefix in _DOCX_XML_PARTS)
        ]
        names.sort(key=lambda name: (0 if name == "word/document.xml" else 1, name))
        for name in names:
            parts.append(_xml_part_text(archive.read(name)))
    return "\n".join(part for part in parts if part.strip())


def _xml_part_text(xml_bytes: bytes) -> str:
    root = ET.fromstring(xml_bytes)
    paragraphs: list[str] = []
    for para in root.iter(f"{W_NS}p"):
        texts: list[str] = []
        for node in para.iter(f"{W_NS}t"):
            if node.text:
                texts.append(node.text)
        line = "".join(texts).strip()
        if line:
            paragraphs.append(line)
    return "\n".join(paragraphs)


# --------------------------------------------------------------------------- DOC
def _extract_doc(path: Path) -> DocumentContent:
    """Legacy .doc: convert with LibreOffice/antiword when present, else scrape strings."""
    warnings: list[str] = []
    header = path.read_bytes()[:8]
    if header.startswith(b"PK\x03\x04"):
        content = _extract_docx(path)
        content.backend = f"{content.backend}+doc-extension"
        content.warnings.extend(warnings)
        return content

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                subprocess.run(
                    [soffice, "--headless", "--convert-to", "docx", "--outdir", tmpdir, str(path)],
                    check=True,
                    capture_output=True,
                    timeout=120,
                )
                converted = next(Path(tmpdir).glob("*.docx"), None)
                if converted:
                    content = _extract_docx(converted)
                    content.backend = "libreoffice+python-docx"
                    return content
        except Exception as exc:
            warnings.append(f"libreoffice conversion failed: {exc.__class__.__name__}")

    antiword = shutil.which("antiword")
    if antiword:
        try:
            result = subprocess.run(
                [antiword, str(path)], check=True, capture_output=True, timeout=60
            )
            return DocumentContent(
                text=result.stdout.decode("utf-8", errors="ignore"), backend="antiword", warnings=warnings
            )
        except Exception as exc:
            warnings.append(f"antiword failed: {exc.__class__.__name__}")

    warnings.append(
        "Legacy .doc parsed with a best-effort text scrape. Install LibreOffice (soffice) for full fidelity."
    )
    return DocumentContent(text=_scrape_binary_strings(path), backend="binary-scrape", warnings=warnings)


def _scrape_binary_strings(path: Path, min_run: int = 4) -> str:
    data = path.read_bytes()
    ascii_runs = re.findall(rb"[\x20-\x7e]{%d,}" % min_run, data)
    utf16_runs = re.findall(rb"(?:[\x20-\x7e]\x00){%d,}" % min_run, data)
    lines = [run.decode("ascii", errors="ignore") for run in ascii_runs]
    lines += [run.decode("utf-16-le", errors="ignore") for run in utf16_runs]
    noise = ("Microsoft", "Word.Document", "MSWordDoc", "Times New Roman", "Root Entry", "WordDocument")
    return "\n".join(line for line in lines if not any(token in line for token in noise))
