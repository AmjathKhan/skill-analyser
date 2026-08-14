"""Render the plain-text sample resumes into real PDF and DOCX files.

The text files under ``data/sample_resumes`` are the source of truth (diff-friendly);
this script produces binary formats so the upload pipeline can be exercised end to end.

    python -m scripts.generate_sample_resumes            # PDF + DOCX
    python -m scripts.generate_sample_resumes --pdf-only
"""

from __future__ import annotations

import argparse
from pathlib import Path

SOURCE_DIR = Path(__file__).resolve().parents[1] / "data" / "sample_resumes"


def render_pdf(text: str, destination: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas

    _width, height = A4
    pdf = canvas.Canvas(str(destination), pagesize=A4)
    pdf.setTitle(destination.stem.replace("_", " ").title())
    left, top, bottom = 18 * mm, height - 20 * mm, 18 * mm
    leading = 12.5
    y = top

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if y <= bottom:
            pdf.showPage()
            y = top
        if not line:
            y -= leading * 0.6
            continue
        is_heading = line.isupper() and len(line) < 60
        pdf.setFont("Helvetica-Bold" if is_heading else "Helvetica", 12 if is_heading else 9.5)
        pdf.drawString(left, y, line[:120])
        y -= leading if not is_heading else leading * 1.3

    pdf.save()


def render_docx(text: str, destination: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    lines = text.splitlines()
    for index, raw_line in enumerate(lines):
        line = raw_line.rstrip()
        if not line:
            continue
        if index == 0:
            document.add_heading(line, level=1)
        elif line.isupper() and len(line) < 60:
            document.add_heading(line.title(), level=2)
        else:
            document.add_paragraph(line)

    document.save(str(destination))


def main() -> int:
    parser = argparse.ArgumentParser(description="Render sample resumes to PDF/DOCX")
    parser.add_argument("--pdf-only", action="store_true")
    parser.add_argument("--docx-only", action="store_true")
    parser.add_argument("--output", default=str(SOURCE_DIR), help="Output directory")
    args = parser.parse_args()

    output_dir = Path(args.output)
    output_dir.mkdir(parents=True, exist_ok=True)

    sources = sorted(SOURCE_DIR.glob("*.txt"))
    if not sources:
        print(f"no .txt sources found in {SOURCE_DIR}")
        return 1

    for source in sources:
        text = source.read_text(encoding="utf-8")
        if not args.docx_only:
            target = output_dir / f"{source.stem}.pdf"
            render_pdf(text, target)
            print(f"wrote {target}")
        if not args.pdf_only:
            target = output_dir / f"{source.stem}.docx"
            render_docx(text, target)
            print(f"wrote {target}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
